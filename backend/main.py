"""
FastAPI Backend — exposes the KPMG chatbot as a REST API
Endpoints:
  POST /chat        ← main chat endpoint (used by Streamlit)
  POST /upload      ← upload document for session-scoped Q&A
  POST /upload/clear← clear uploaded docs for a session
  GET  /health      ← health check (used by Azure App Service)
  GET  /tickets     ← view all helpdesk tickets (admin)
  GET  /logs        ← view all RAID logs (admin)
  GET  /approvals   ← view all approvals (admin)
  GET  /docs        ← Swagger UI (auto-generated)

Validated with Pydantic. Tested via Swagger UI + Postman.
"""
import logging
import uuid
import os
import sys
from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from dotenv import load_dotenv

from router.graph import run_chat
from agents.workflow_agent import get_all_tickets, get_all_logs, get_all_approval_records
from ingestion.vector_store import index_exists
from ingestion.upload_handler import (
    ingest_uploaded_file, get_session_files, clear_session, get_session_stats
)
from run_ingestion import run_ingestion
from utils.watchdog_service import start_watchdog

load_dotenv()

# ── Structured logging ───────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s — %(message)s"
)
logger = logging.getLogger(__name__)


# ── FastAPI app ──────────────────────────────────────
app = FastAPI(
    title="KPMG PMO AI Chatbot",
    description=(
        "Multi-agent RAG chatbot for KPMG PMO team. "
        "Supports document Q&A, KPI retrieval, IT helpdesk tickets, and workflow automation. "
        "Built with LangGraph + Groq (LLaMA 3.1) + FAISS + HuggingFace embeddings."
    ),
    version="2.0.0",
    docs_url="/docs",
    redoc_url="/redoc",
)

# Allow Streamlit frontend to call this API
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],   # tighten in production
    allow_methods=["*"],
    allow_headers=["*"],
)


# ── Lifecycle ────────────────────────────────────────
@app.on_event("startup")
def startup_event():
    """Run RAG bootstrap and start watchdog on startup."""
    logger.info("Initializing RAG Engine...")
    # 1. Recursive bootstrap walk of /local_storage
    run_ingestion(folder="local_storage")
    
    # 2. Start watchdog for real-time incremental updates
    app.state.observer = start_watchdog()
    logger.info("RAG Engine & Watchdog initialized.")


@app.on_event("shutdown")
def shutdown_event():
    """Stop watchdog on shutdown."""
    if hasattr(app.state, "observer"):
        app.state.observer.stop()
        app.state.observer.join()
        logger.info("Watchdog stopped.")


# ── Pydantic models ──────────────────────────────────
class ChatRequest(BaseModel):
    query: str = Field(
        ...,
        min_length=1,
        max_length=15000,
        description="User's natural language question or command",
        example="What are the delivery risks mentioned in the Q3 report?"
    )
    session_id: str = Field(
        default=None,
        description="Session ID for accessing uploaded documents (optional)"
    )
    upload_only: bool = Field(
        default=False,
        description="If True, only search the session's uploaded documents."
    )
    agent_override: str = Field(
        default="auto",
        description="Explicitly route to a specific agent (auto, docs, kpis, helpdesk, automation)"
    )
    chat_history: list = Field(
        default=[],
        description="Last N messages for conversation context [{role, content}]"
    )


class ChatResponse(BaseModel):
    answer:     str  = Field(..., description="AI-generated response")
    agent_used: str  = Field(..., description="Which agent handled this query")
    sources:    list = Field(..., description="Source documents or systems used")
    status:     str  = Field(default="success")


class HealthResponse(BaseModel):
    status:       str
    index_loaded: bool
    version:      str
    agents:       list


# ── Routes ───────────────────────────────────────────

@app.get("/health", response_model=HealthResponse, tags=["System"])
def health_check():
    """Health check — verifies the app and all agents are running."""
    return {
        "status":       "healthy",
        "index_loaded": index_exists(),
        "version":      "2.0.0",
        "agents":       ["Retrieval Agent", "API Agent", "Helpdesk Agent", "Workflow Agent"]
    }


@app.post("/chat", response_model=ChatResponse, tags=["Chat"])
def chat(request: ChatRequest):
    """
    Main chat endpoint.
    Routes query through LangGraph → Retrieval / API / Helpdesk / Workflow agent.
    Returns grounded answer with source attribution.
    """
    logger.info(f"[/chat] Query: {request.query[:80]}...")

    try:
        result = run_chat(
            request.query, 
            session_id=request.session_id, 
            upload_only=request.upload_only,
            agent_override=request.agent_override,
            chat_history=request.chat_history
        )
        logger.info(f"[/chat] Agent: {result['agent_used']}")

        return ChatResponse(
            answer     = result["answer"],
            agent_used = result["agent_used"],
            sources    = result["sources"],
            status     = "success"
        )

    except ValueError as e:
        logger.error(f"[/chat] Config error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

    except Exception as e:
        logger.error(f"[/chat] Unexpected error: {e}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Internal error: {str(e)}"
        )


@app.get("/tickets", tags=["Admin"])
def get_tickets():
    """View all IT helpdesk tickets (from ServiceNow or mock)."""
    tickets = get_all_tickets()
    return {"count": len(tickets), "tickets": tickets}


@app.get("/logs", tags=["Admin"])
def get_logs():
    """View all RAID log entries (from SharePoint or mock)."""
    logs = get_all_logs()
    return {"count": len(logs), "logs": logs}


@app.get("/logs/system", tags=["Admin"])
def get_system_logs():
    """View all system logs (from System_Log.csv)."""
    from connectors.azure_insights_connector import query_logs
    logs = query_logs(limit=50)
    return {"count": len(logs), "logs": logs}


@app.get("/approvals", tags=["Admin"])
def get_approvals():
    """View all approval requests (from Teams or mock)."""
    approvals = get_all_approval_records()
    return {"count": len(approvals), "approvals": approvals}


@app.get("/onboarding/{emp_id}", tags=["Admin"])
def get_onboarding_details(emp_id: str):
    """Fetch employee details by EMP ID from the onboarding tracker."""
    from connectors.outlook_connector import _get_record_by_emp_id
    record = _get_record_by_emp_id(emp_id)
    if not record:
        raise HTTPException(status_code=404, detail=f"No active record found for {emp_id}")
    return record


@app.get("/download/tracker/{tracker_type}", tags=["Admin"])
def download_tracker(tracker_type: str):
    """Download the onboarding, offboarding, tagging, tickets, or raid Excel tracker."""
    tracker_map = {
        "onboarding": "Onboarding_Records.xlsx",
        "offboarding": "Offboarding_Records.xlsx",
        "tagging": "Resource_Tagging_Records.xlsx",
        "tickets": "IT_Helpdesk_Tickets.xlsx",
        "raid": "RAID_Log_KPMG.csv"
    }
    filename = tracker_map.get(tracker_type.lower())
    if not filename:
        raise HTTPException(status_code=400, detail="Invalid tracker type")
        
    file_path = Path("local_storage") / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail=f"Tracker file not found: {filename}")
        
    media_type = "text/csv" if filename.endswith(".csv") else "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        
    return StreamingResponse(
        FileResponse(
            path=file_path, 
            filename=filename, 
            media_type=media_type
        )
    )

@app.get("/download/logs/txt", tags=["Admin"])
def download_system_logs_txt():
    """Download system logs as a plain text file."""
    from connectors.azure_insights_connector import query_logs
    try:
        logs = query_logs(hours=48, limit=100)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
        
    if not logs:
        raise HTTPException(status_code=404, detail="No system logs found.")
        
    content = ["KPMG System Logs Export", "=" * 50, ""]
    for log in logs:
        time = log.get('timestamp', '')[:19].replace('T', ' ')
        content.append(f"[{time}] {log.get('severity', 'INFO').upper()} | Source: {log.get('source', 'Unknown')}")
        content.append(f"Message: {log.get('message', '')}")
        content.append("-" * 50)
        
    from fastapi.responses import StreamingResponse
    import io
    
    stream = io.StringIO("\n".join(content))
    response = StreamingResponse(iter([stream.getvalue()]), media_type="text/plain")
    response.headers["Content-Disposition"] = "attachment; filename=system_logs_export.txt"
    return response


@app.get("/", tags=["System"])
def root():
    """Root endpoint — confirms API is running."""
    return {
        "message": "KPMG PMO AI Chatbot API v2.0 is running.",
        "docs":    "/docs",
        "health":  "/health",
        "chat":    "POST /chat",
        "upload":  "POST /upload",
        "agents":  ["Retrieval", "API", "Helpdesk", "Workflow"]
    }


# ── Upload endpoints ─────────────────────────────────

ALLOWED_EXTENSIONS = {"pdf", "docx", "doc", "xlsx", "xls", "csv", "txt", "md",
                       "png", "jpg", "jpeg", "gif", "pptx"}
MAX_FILE_SIZE = 20 * 1024 * 1024   # 20 MB


@app.post("/upload", tags=["Upload"])
async def upload_file(
    file: UploadFile = File(..., description="Document to upload"),
    session_id: str = Form(..., description="Session ID for the upload")
):
    """
    Upload a document for session-scoped Q&A.
    The file is parsed, chunked, embedded, and stored in a temporary FAISS index.
    It will be searchable via /chat when the same session_id is provided.
    """
    # Validate extension
    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(400, f"Unsupported file type: .{ext}. Allowed: {', '.join(ALLOWED_EXTENSIONS)}")

    # Read file
    contents = await file.read()
    if len(contents) > MAX_FILE_SIZE:
        raise HTTPException(400, f"File too large ({len(contents) // 1024 // 1024}MB). Max: 20MB.")

    if len(contents) == 0:
        raise HTTPException(400, "Empty file uploaded.")

    logger.info(f"[/upload] {file.filename} ({len(contents)} bytes) → session {session_id[:8]}...")

    try:
        result = ingest_uploaded_file(session_id, file.filename, contents)
        return result
    except Exception as e:
        logger.error(f"[/upload] Failed: {e}", exc_info=True)
        raise HTTPException(500, f"Upload processing failed: {str(e)}")


@app.post("/upload/clear", tags=["Upload"])
def clear_uploads(session_id: str = Form(...)):
    """Clear all uploaded documents for a session."""
    clear_session(session_id)
    return {"status": "cleared", "session_id": session_id}


@app.get("/upload/files", tags=["Upload"])
def list_uploaded_files(session_id: str):
    """List uploaded files for a session."""
    stats = get_session_stats(session_id)
    return stats


# ── Document conversion endpoints ────────────────────
import re
import io
from pathlib import Path
from fastapi.responses import StreamingResponse
from run_ingestion import run_ingestion

DOC_BASE_PATH = Path(__file__).parent.parent / "local_storage" / "Project_Flowchart"


def _html_to_text(html_content: str) -> str:
    """Strip HTML tags to get plain text."""
    text = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL)
    text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL)
    text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'</(p|div|h[1-6]|li|tr)>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'<[^>]+>', '', text)
    text = re.sub(r'&nbsp;', ' ', text)
    text = re.sub(r'&amp;', '&', text)
    text = re.sub(r'&lt;', '<', text)
    text = re.sub(r'&gt;', '>', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()




@app.get("/convert/{filename}", tags=["Documents"])
def convert_document(filename: str, format: str = "pdf"):
    """
    Convert an HTML vault document to PDF, DOCX, or XLSX.
    """
    full_path = DOC_BASE_PATH / filename
    if not full_path.exists():
        raise HTTPException(404, f"Document '{filename}' not found.")

    html_content = full_path.read_text(encoding="utf-8", errors="ignore")
    plain_text = _html_to_text(html_content)
    base_name = filename.rsplit(".", 1)[0]

    if format == "pdf":
        try:
            try:
                from fpdf import FPDF
            except ImportError:
                import subprocess
                import sys
                try:
                    # Auto-install within the exact backend environment
                    subprocess.check_call([sys.executable, "-m", "pip", "install", "fpdf2", "python-docx", "openpyxl"])
                    from fpdf import FPDF
                except Exception as e:
                    raise HTTPException(500, f"Auto-install failed: {e}. Env: {sys.executable}")
            
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.set_font("helvetica", size=10)
            
            # fpdf2 with default fonts only supports Latin-1. 
            # We strip out non-Latin-1 characters to avoid UnicodeEncodeError.
            safe_text = plain_text.encode('latin-1', 'replace').decode('latin-1')
            
            # pdf.write automatically wraps text instead of crashing on long un-spaced strings
            pdf.write(6, safe_text)
            buf = io.BytesIO(pdf.output())
            buf.seek(0)
            return StreamingResponse(buf, media_type="application/pdf",
                                     headers={"Content-Disposition": f"attachment; filename={base_name}.pdf"})
        except Exception as e:
            import traceback
            error_msg = traceback.format_exc()
            with open("/Users/shouvikpaul/Downloads/kpmg-chatbot/backend_pdf_error.log", "w") as f:
                f.write(error_msg)
            raise HTTPException(500, f"PDF generation failed. See backend_pdf_error.log.")


    elif format == "docx":
        try:
            from docx import Document
        except ImportError:
            raise HTTPException(500, "python-docx not installed. Run: pip install python-docx")
        doc = Document()
        doc.add_heading(base_name, level=1)
        for para in plain_text.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                 headers={"Content-Disposition": f"attachment; filename={base_name}.docx"})

    elif format == "xlsx":
        try:
            from openpyxl import Workbook
        except ImportError:
            raise HTTPException(500, "openpyxl not installed. Run: pip install openpyxl")
        wb = Workbook()
        ws = wb.active
        ws.title = base_name[:31]
        ws.append(["Content from: " + filename])
        ws.append([])
        for i, line in enumerate(plain_text.split("\n"), 1):
            if line.strip():
                ws.append([line.strip()])
        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                 headers={"Content-Disposition": f"attachment; filename={base_name}.xlsx"})

    else:
        raise HTTPException(400, f"Unsupported format: {format}. Use pdf, docx, or xlsx.")
